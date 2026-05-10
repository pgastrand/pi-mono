#!/usr/bin/env python3
"""Generate codebase documentation DOCX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text):
    p = doc.add_paragraph(text)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

doc = Document()

# Title
title = doc.add_heading('Pi Monorepo Codebase Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Version 0.0.3 | MIT License')
doc.add_paragraph()

# Overview
heading(doc, 'Overview', 1)
para(doc, 'Pi (shittycodingagent.ai) is a TypeScript monorepo for building AI coding agents and LLM tooling. It provides a complete ecosystem for developing, deploying, and interacting with large language model-based agents.')
para(doc, 'The project consists of 7 interconnected packages that form a complete AI development environment.')

# Packages Table
heading(doc, 'Packages', 1)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Package'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'NPM Package'

packages = [
    ('pi-ai', 'Unified multi-provider LLM API', '@mariozechner/pi-ai'),
    ('pi-agent-core', 'Stateful agent runtime with tool execution', '@mariozechner/pi-agent-core'),
    ('pi-coding-agent', 'Interactive CLI coding agent', '@mariozechner/pi-coding-agent'),
    ('pi-tui', 'Terminal UI framework with differential rendering', '@mariozechner/pi-tui'),
    ('pi-web-ui', 'Web components for AI chat interfaces', '@mariozechner/pi-web-ui'),
    ('pi-mom', 'Slack bot with self-managing capabilities', '@mariozechner/pi-mom'),
    ('pi-pods', 'CLI for vLLM GPU pod deployment', '@mariozechner/pi-pods'),
]

for i, (pkg, desc, npm) in enumerate(packages):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = pkg
    row_cells[1].text = desc
    row_cells[2].text = npm

doc.add_paragraph()

# Architecture
heading(doc, 'Architecture', 1)
para(doc, 'The codebase follows a layered architecture where higher-level packages build upon lower-level ones:')

heading(doc, '1. pi-ai (Foundation Layer)', 2)
para(doc, 'The core LLM abstraction layer providing:')
bullet(doc, 'Unified API for 15+ LLM providers (OpenAI, Anthropic, Google, etc.)')
bullet(doc, 'Streaming and non-streaming request handling')
bullet(doc, 'Tool calling (function calling) with TypeBox schema validation')
bullet(doc, 'Token usage and cost tracking')
bullet(doc, 'Cross-provider context handoffs')
bullet(doc, 'OAuth authentication support')

heading(doc, '2. pi-agent-core (Agent Runtime)', 2)
para(doc, 'Stateful agent implementation on top of pi-ai:')
bullet(doc, 'Event-driven architecture with subscribe/publish pattern')
bullet(doc, 'Tool execution with preflight checks and hooks')
bullet(doc, 'Message transformation and conversion pipeline')
bullet(doc, 'Steering and follow-up message queues')
bullet(doc, 'Parallel and sequential tool execution modes')

heading(doc, '3. pi-tui (Terminal UI)', 2)
para(doc, 'Minimal terminal UI framework:')
bullet(doc, 'Differential rendering (only updates changed lines)')
bullet(doc, 'Synchronized output using CSI 2026 for flicker-free updates')
bullet(doc, 'Component-based architecture')
bullet(doc, 'Support for Editor, SelectList, Markdown, Image components')

heading(doc, '4. pi-coding-agent (CLI Application)', 2)
para(doc, 'The main interactive coding agent:')
bullet(doc, 'Built on pi-tui for terminal interface')
bullet(doc, 'File system tools (read, write, edit, bash)')
bullet(doc, 'Session management with branching and compaction')
bullet(doc, 'Extensions system for customization')
bullet(doc, 'Skills and prompt templates')
bullet(doc, 'Theme support')

heading(doc, '5. pi-mom (Slack Bot)', 2)
para(doc, 'Slack integration with autonomous capabilities:')
bullet(doc, 'Socket Mode for real-time message handling')
bullet(doc, 'Self-managing tool installation')
bullet(doc, 'Per-channel workspace isolation')
bullet(doc, 'Memory across sessions via MEMORY.md files')
bullet(doc, 'Events system for scheduled tasks')
bullet(doc, 'Docker sandbox execution')

heading(doc, '6. pi-pods (Deployment Tool)', 2)
para(doc, 'GPU pod management for vLLM:')
bullet(doc, 'Automated vLLM installation and configuration')
bullet(doc, 'Predefined model configurations (Qwen, GLM, GPT-OSS)')
bullet(doc, 'Multi-GPU allocation and tensor parallelism')
bullet(doc, 'OpenAI-compatible API endpoints')
bullet(doc, 'Interactive agent for testing')

heading(doc, '7. pi-web-ui (Web Components)', 2)
para(doc, 'Browser-based chat interface:')
bullet(doc, 'Built with mini-lit web components')
bullet(doc, 'Tailwind CSS v4 styling')
bullet(doc, 'IndexedDB storage for sessions')
bullet(doc, 'Artifacts panel for HTML/SVG/Markdown')
bullet(doc, 'Attachment handling (PDF, DOCX, images)')

# Key Technologies
heading(doc, 'Key Technologies', 1)
bullet(doc, 'Runtime: Node.js 20+')
bullet(doc, 'Language: TypeScript')
bullet(doc, 'Linting: Biome')
bullet(doc, 'Testing: Vitest')
bullet(doc, 'Package Manager: npm workspaces')
bullet(doc, 'Schema Validation: TypeBox + AJV')

# Design Philosophy
heading(doc, 'Design Philosophy', 1)
para(doc, 'Pi follows a deliberately constrained approach:')

heading(doc, 'What Pi Does NOT Include', 2)
bullet(doc, 'No MCP - Uses CLI tools via skills (markdown READMEs)')
bullet(doc, 'No sub-agents - Spawn via tmux or build with extensions')
bullet(doc, 'No permission popups - Use containers or build with extensions')
bullet(doc, 'No plan mode - Write plans to files or build with extensions')
bullet(doc, 'No built-in to-dos - Use TODO.md or build with extensions')

heading(doc, 'Design Principles', 2)
bullet(doc, 'Aggressively extensible via extensions, skills, themes, prompt templates')
bullet(doc, 'Minimal core with maximum customization potential')
bullet(doc, 'Build what you need rather than bloated features')
bullet(doc, 'Package ecosystem for sharing via npm/git')

# File Structure
heading(doc, 'File Structure', 1)
code_block(doc, '''pi-monorepo/
├── packages/
│   ├── agent/           # pi-agent-core
│   ├── ai/              # pi-ai (LLM APIs)
│   ├── coding-agent/    # Main CLI application
│   ├── tui/             # Terminal UI framework
│   ├── web-ui/          # Web components
│   ├── mom/             # Slack bot
│   └── pods/            # GPU pod deployment
├── scripts/             # Build and utility scripts
├── .github/workflows/   # CI/CD pipelines
└── tsconfig.base.json   # Shared TypeScript config''')

# Development
heading(doc, 'Development Setup', 1)
para(doc, 'Getting started:')
code_block(doc, '''npm install          # Install all dependencies
npm run build        # Build all packages
npm run check        # Lint and type check
./test.sh            # Run tests''')

# Contributing
heading(doc, 'Contributing', 1)
para(doc, 'See CONTRIBUTING.md for guidelines and AGENTS.md for project-specific rules for both humans and agents.')

# License
heading(doc, 'License', 1)
para(doc, 'MIT License - See LICENSE file for details.')

# Save
os.makedirs('output/doc', exist_ok=True)
doc.save('output/doc/doc.docx')
print('Created: output/doc/doc.docx')